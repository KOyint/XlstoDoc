import openpyxl
import docx
import datetime
import re 

#modify the default date time format to dd/mm/yyyy
def Datemod(dt):
	return str(datetime.datetime.strptime(dt,'%Y-%m-%d' ).strftime('%d/%m/%Y'))

#method to save new document?	
def SaveDoc(document,name):
	document.save(name +'.docx')

def ClearAndAdd(run , value):
	run.clear()
	run.add_text(str(value))


#load template
Template = docx.Document('AS.docx')

#load data excelsheet 
Data = openpyxl.load_workbook('VO.xlsx').active

#define the col for each value in date excel
PackageCol = 1
PriceCol = 2
CNameCol = 3
SDateCol = 4
EDateCol = 5
DurCol = 6
DepoCol = 7
ExtCol = 8
NameCol = 9
ContactCol = 10
EmailCol = 11
MeetCol = 12

#para & run value for doc editing (ownself found out using py docx)

CNameRun = Template.paragraphs[10].runs[13]
PackageRun = Template.paragraphs[16].runs[9]
RentRun = Template.paragraphs[18].runs[14]
DepoRun = Template.paragraphs[19].runs[17]
DurRun = Template.paragraphs[21].runs[13]
MeetRun = Template.paragraphs[23].runs[28]
ContactRun = Template.paragraphs[28].runs[8]


for i in range(1,2):
	Package = str(Data.cell(i,PackageCol).value)
	CName = str(Data.cell(i,CNameCol).value)
	Rent = str(Data.cell(i,PriceCol).value)
	Depo = str(Data.cell(i,DepoCol).value)
	Start = Datemod(str(Data.cell(i,SDateCol).value.date()))
	End	= Datemod(str(Data.cell(i,EDateCol).value.date()))
	Dur = str(Start + ' - ' + End)
	Meet = str(Data.cell(i,MeetCol).value)
	Contact = str(Data.cell(i,ContactCol).value)
	
	ClearAndAdd(CNameRun,CName)
	ClearAndAdd(PackageRun,Package)
	ClearAndAdd(RentRun,"$"+Rent)
	ClearAndAdd(DepoRun,"$"+Depo)
	ClearAndAdd(DurRun,Dur)
	ClearAndAdd(MeetRun,Meet)
	ClearAndAdd(ContactRun,Contact)
	
	SaveDoc(Template,CName)
	

	




	
